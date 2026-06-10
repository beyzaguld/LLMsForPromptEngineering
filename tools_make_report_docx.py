"""
Convert FINAL_REPORT.md into a formatted Word document (FINAL_REPORT.docx).

Handles: H1/H2/H3 headings, paragraphs with inline **bold**/`code`, bullet lists,
fenced code blocks, blockquotes, and GitHub-style pipe tables.

Usage:  python3 tools_make_report_docx.py
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "FINAL_REPORT.md"
OUT = "FINAL_REPORT.docx"

CODE_GRAY = RGBColor(0x24, 0x29, 0x2E)
MONO = "Consolas"


def add_inline(paragraph, text):
    """Render inline **bold** and `code` spans into runs."""
    for part in re.split(r"(\*\*.*?\*\*|`[^`]*`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = MONO
            r.font.size = Pt(9.5); r.font.color.rgb = CODE_GRAY
        else:
            paragraph.add_run(part)


def build():
    doc = Document()
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)

    lines = open(SRC, encoding="utf-8").read().split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # fenced code block
        if line.strip().startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1  # closing fence
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run("\n".join(buf))
            r.font.name = MONO; r.font.size = Pt(9)
            r.font.color.rgb = CODE_GRAY
            continue

        # table block
        if line.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[ :\-|]+\|\s*$", lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.paragraphs[0].text = ""
                add_inline(cell.paragraphs[0], h)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for j, val in enumerate(row):
                    if j < len(cells):
                        cells[j].paragraphs[0].text = ""
                        add_inline(cells[j].paragraphs[0], val)
            doc.add_paragraph()
            continue

        # headings
        if line.startswith("# "):
            h = doc.add_heading(level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(h, line[2:].strip())
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        # horizontal rule
        elif line.strip() == "---":
            pass
        # blockquote
        elif line.strip().startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, line.strip().lstrip(">").strip())
        # bullet list
        elif re.match(r"^\s*[-*] ", line):
            indent = len(line) - len(line.lstrip())
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_inline(p, re.sub(r"^\s*[-*] ", "", line))
        # numbered list
        elif re.match(r"^\s*\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\s*\d+\. ", "", line))
        # blank
        elif line.strip() == "":
            pass
        # normal paragraph (gather following non-empty, non-special lines)
        else:
            buf = [line]
            while (i + 1 < n and lines[i + 1].strip() != ""
                   and not re.match(r"^(#|\||```|>|\s*[-*] |\s*\d+\. |---)", lines[i + 1])):
                i += 1
                buf.append(lines[i])
            p = doc.add_paragraph()
            add_inline(p, " ".join(s.strip() for s in buf))
        i += 1

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
